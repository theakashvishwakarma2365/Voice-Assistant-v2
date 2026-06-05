from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Color palette ─────────────────────────────────────────────────────────────
C_DARK   = RGBColor(0x1A, 0x1A, 0x2E)   # deep navy
C_BLUE   = RGBColor(0x16, 0x21, 0x3E)   # section header bg
C_ACCENT = RGBColor(0x0F, 0x3C, 0x78)   # accent blue
C_LIGHT  = RGBColor(0xE8, 0xF0, 0xFE)   # light blue tint
C_GREEN  = RGBColor(0x1B, 0x87, 0x54)   # green highlight
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_GRAY   = RGBColor(0x64, 0x74, 0x87)
C_TABLE_H= RGBColor(0x0F, 0x3C, 0x78)
C_TABLE_A= RGBColor(0xF0, 0xF4, 0xFF)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)

def set_cell_borders(cell, color='BBBBBB'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def para_space(doc, before=0, after=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    return p

def add_title_block(doc):
    # Cover banner
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run('VoiceDesk')
    run.font.size    = Pt(36)
    run.font.bold    = True
    run.font.color.rgb = C_ACCENT

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run('System Architecture Document')
    r2.font.size = Pt(18)
    r2.font.color.rgb = C_DARK

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(30)
    r3 = p3.add_run('Handheld Voice Productivity Assistant')
    r3.font.size = Pt(13)
    r3.font.italic = True
    r3.font.color.rgb = C_GRAY

    # Meta table
    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta = [
        ('Project',  'VoiceDesk v1.0'),
        ('Author',   'System Design Session'),
        ('Date',     'June 2026'),
        ('Status',   'Architecture Approved — Ready for Implementation'),
    ]
    for i,(k,v) in enumerate(meta):
        kc = tbl.rows[i].cells[0]
        vc = tbl.rows[i].cells[1]
        set_cell_bg(kc, C_TABLE_H)
        set_cell_bg(vc, C_TABLE_A)
        kr = kc.paragraphs[0].add_run(k)
        kr.font.bold = True; kr.font.color.rgb = C_WHITE; kr.font.size = Pt(10)
        vr = vc.paragraphs[0].add_run(v)
        vr.font.size = Pt(10); vr.font.color.rgb = C_DARK
    para_space(doc, after=10)

def section_heading(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f'{number}  {title}')
    run.font.size  = Pt(14)
    run.font.bold  = True
    run.font.color.rgb = C_ACCENT
    # underline rule
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:color'), '0F3C78')
    pBdr.append(bottom)
    pPr.append(pBdr)

def sub_heading(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(title)
    run.font.size  = Pt(11)
    run.font.bold  = True
    run.font.color.rgb = C_DARK

def body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = C_DARK
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.5 + level*0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = C_DARK

def code_block(doc, code_text):
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
        # light grey bg per paragraph
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F3F4F6')
        pPr.append(shd)
    para_space(doc, after=4)

def make_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1+len(rows), cols=n_cols)
    tbl.style = 'Table Grid'
    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, C_TABLE_H)
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.color.rgb = C_WHITE
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for ri, row in enumerate(rows):
        bg = C_TABLE_A if ri % 2 == 0 else C_WHITE
        for ci, val in enumerate(row):
            cell = tbl.rows[ri+1].cells[ci]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = C_DARK
    # Column widths
    if col_widths:
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[ci])
    para_space(doc, after=6)
    return tbl

def note_box(doc, text, color=None):
    color = color or C_LIGHT
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f'  ℹ  {text}')
    run.font.size   = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = C_ACCENT
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'E8F0FE')
    pPr.append(shd)

# ═══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════

add_title_block(doc)
doc.add_page_break()

# ── TABLE OF CONTENTS ─────────────────────────────────────────────────────────
section_heading(doc, '', 'Table of Contents')
toc = [
    '1  Executive Summary',
    '2  System Overview',
    '3  Hardware Architecture',
    '4  Software Architecture',
    '5  Data Architecture',
    '6  Network & Communication',
    '7  UI / UX Design',
    '8  Voice Command Flow',
    '9  Google Sheets Integration',
    '10  Raspberry Pi 5 Server Setup',
    '11  Complete Wiring Reference',
    '12  Bill of Materials',
    '13  Development Roadmap',
    '14  Architectural Decisions & Trade-offs',
]
for item in toc:
    bullet(doc, item)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '1', 'Executive Summary')
body(doc, (
    'VoiceDesk is a local-first, voice-controlled productivity assistant built on a two-tier '
    'hardware architecture. A handheld ESP32 device serves as the user interface — capturing '
    'voice input, displaying status on a 3.5-inch TFT screen, and providing physical button '
    'navigation. A Raspberry Pi 5 acts as the local AI server — running Whisper speech-to-text, '
    'an Ollama language model for intent parsing, Piper for text-to-speech synthesis, and a '
    'FastAPI backend that syncs all data to Google Sheets.'
))
body(doc, (
    'All AI processing is local. No user data leaves the home network. The system responds to '
    'natural language voice commands for task management, meeting scheduling, project tracking, '
    'Pomodoro sessions, and reminders — with intelligent follow-up questions and context-aware '
    'suggestions powered by the on-device LLM.'
))

make_table(doc,
    ['Attribute', 'Value'],
    [
        ('Primary Interface',  'Voice commands (always-on wake word or PTT button)'),
        ('Display',            '3.5" ILI9486 SPI TFT, 480×320, button-navigated'),
        ('Handheld MCU',       'ESP32 DevKit v1 (ESP-WROOM-32)'),
        ('AI Server',          'Raspberry Pi 5 (4GB RAM)'),
        ('Speech-to-Text',     'faster-whisper (base.en model, local)'),
        ('Language Model',     'Ollama phi3:mini (3.8B parameters, local)'),
        ('Text-to-Speech',     'Piper TTS (offline, natural voice)'),
        ('Data Storage',       'SQLite (local) + Google Sheets (cloud sync)'),
        ('Connectivity',       'WiFi 802.11n — ESP32 to Pi5 on local LAN'),
        ('Power (handheld)',   '3.7V Li-Po 3000mAh, ~6–8 hour runtime'),
    ],
    col_widths=[2.2, 4.0]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 2. SYSTEM OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '2', 'System Overview')
body(doc, 'The system consists of two physical nodes connected over a local WiFi network:')

sub_heading(doc, '2.1  Two-Tier Architecture')
make_table(doc,
    ['Tier', 'Device', 'Role', 'Processing'],
    [
        ('Tier 1 — Client', 'ESP32 Handheld', 'Voice capture, display, button input, audio playback', 'Thin client — wake word only'),
        ('Tier 2 — Server', 'Raspberry Pi 5', 'STT, LLM inference, TTS, DB, Sheets sync', 'All heavy AI workloads'),
    ],
    col_widths=[1.5, 1.5, 2.5, 1.7]
)

sub_heading(doc, '2.2  High-Level Data Flow')
code_block(doc, """
[User speaks]
     │
     ▼
[ESP32] Wake word detected (Porcupine, on-device)
     │  Hold MIC button (PTT mode)
     │
     ▼ WebSocket — PCM audio stream (16kHz)
[Pi 5 FastAPI Server]
     │
     ├─► [faster-whisper]  →  transcript text
     │
     ├─► [Ollama phi3:mini]  →  intent + entities JSON
     │         │
     │    SCHEDULE_MEETING | ADD_TASK | QUERY | DELETE |
     │    NEW_SHEET | POMODORO | REMINDER | SWITCH_SHEET
     │
     ├─► [Action Executor]  →  SQLite write  →  Google Sheets sync
     │
     ├─► [Response Generator / Ollama]  →  natural language reply
     │
     └─► [Piper TTS]  →  WAV audio
             │
             ▼ WebSocket — WAV + JSON display payload
[ESP32]  plays audio + updates screen
""")

note_box(doc, 'All processing is local. Internet is only needed for Google Sheets sync. The device works fully offline — data queues and syncs when connectivity is restored.')

# ═══════════════════════════════════════════════════════════════════════════════
# 3. HARDWARE ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '3', 'Hardware Architecture')

sub_heading(doc, '3.1  ESP32 Handheld Device — Component Roles')
make_table(doc,
    ['Component', 'Model', 'Interface', 'Role'],
    [
        ('Microcontroller',  'ESP32-WROOM-32 DevKit v1', '— (main board)', 'Central processor, WiFi, I2S, SPI'),
        ('Display',          'MPI3501 3.5" ILI9486 480×320', 'SPI', 'User interface, face animation, menus'),
        ('Microphone',       'INMP441 MEMS I2S', 'I2S (port 0)', 'Capture voice audio at 16kHz'),
        ('Speaker Amp',      'MAX98357A Class-D 3W', 'I2S (port 1)', 'Drive 8Ω speaker for TTS playback'),
        ('Speaker',          '40mm 8Ω 3W', 'Analog out', 'Audio output'),
        ('Buttons ×5',       '12×12mm tactile momentary', 'GPIO (digital in)', 'BACK / UP / DOWN / OK / MIC-PTT'),
        ('Battery',          '3.7V Li-Po 3000mAh', 'Power rail', '6–8 hour runtime'),
        ('Charge module',    'TP4056 USB-C + MT3608 5V boost', 'Power rail', 'Charging + 5V regulation'),
        ('Power switch',     'SPDT slide', 'Power rail', 'Hard on/off'),
        ('LED indicators',   '3mm LED ×2 (green + blue)', 'GPIO', 'Charging status / power on'),
        ('Capacitors',       '100µF + 10µF electrolytic', 'Power rail', 'Audio noise decoupling'),
    ],
    col_widths=[1.4, 2.0, 1.2, 2.6]
)

sub_heading(doc, '3.2  Button Function Map')
make_table(doc,
    ['Button', 'GPIO', 'Short Press', 'Long Press'],
    [
        ('◄ BACK',    'GPIO 34', 'Go back / cancel',        'Return to home screen'),
        ('▲ UP',      'GPIO 35', 'Scroll up / previous',    '—'),
        ('▼ DOWN',    'GPIO 33', 'Scroll down / next',      '—'),
        ('✓ OK',      'GPIO 13', 'Select / confirm',        'Open context menu'),
        ('🎤 MIC/PTT','GPIO 14', 'Hold to record (PTT)',    'Toggle always-listen mode'),
    ],
    col_widths=[1.0, 0.8, 2.2, 2.2]
)
note_box(doc, 'GPIO 34 and 35 are input-only pins on ESP32 — perfect for buttons. All buttons wired with INPUT_PULLUP (active LOW on press).')

sub_heading(doc, '3.3  Raspberry Pi 5 Server')
make_table(doc,
    ['Component', 'Spec', 'Purpose'],
    [
        ('Raspberry Pi 5',   '4GB RAM, quad-core Cortex-A76 @ 2.4GHz', 'Main AI inference server'),
        ('NVMe SSD',         '256GB M.2 2242 via official M.2 HAT+',   'OS + models + database (10× faster than SD)'),
        ('Active Cooler',    'Official Pi5 fan + heatsink',             'Prevent CPU throttle under AI load'),
        ('Power Supply',     '27W USB-C PD (official)',                 'Stable power for sustained workloads'),
        ('Network',          'Ethernet (preferred) or onboard WiFi',    'LAN connection to ESP32 device'),
    ],
    col_widths=[1.6, 2.8, 2.0]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 4. SOFTWARE ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '4', 'Software Architecture')

sub_heading(doc, '4.1  ESP32 Firmware Stack')
make_table(doc,
    ['Layer', 'Technology', 'Notes'],
    [
        ('Framework',       'Arduino (ESP-IDF via Arduino core)',    'Familiar, large library ecosystem'),
        ('Display driver',  'TFT_eSPI library',                     'ILI9486 config, SPI @ 27MHz'),
        ('UI renderer',     'Custom C++ sprite engine',             'Double-buffer sprites, face animation'),
        ('Audio capture',   'ESP32 I2S driver (built-in)',          'I2S0, 16kHz, 32-bit → trim to 16-bit'),
        ('Audio playback',  'ESP32 I2S driver (built-in)',          'I2S1, PCM WAV streaming'),
        ('Wake word',       'Porcupine SDK (Picovoice, free tier)', '~30KB model, runs fully on-chip'),
        ('Networking',      'WiFi + WebSocket (ArduinoWebsockets)', 'Persistent WS connection to Pi5'),
        ('Config storage',  'NVS (Non-Volatile Storage)',           'WiFi credentials, server IP'),
        ('OTA updates',     'ArduinoOTA / ElegantOTA',              'Update firmware over WiFi'),
    ],
    col_widths=[1.5, 2.2, 2.5]
)

sub_heading(doc, '4.2  Raspberry Pi 5 Server Stack')
make_table(doc,
    ['Layer', 'Technology', 'Notes'],
    [
        ('OS',              'Raspberry Pi OS Lite 64-bit',          'Headless, no desktop needed'),
        ('Runtime',         'Python 3.11',                          'Main application language'),
        ('API server',      'FastAPI + Uvicorn',                    'WebSocket + REST endpoints'),
        ('STT engine',      'faster-whisper (base.en)',             '~74MB model, ~1.5s latency on Pi5'),
        ('LLM',             'Ollama + phi3:mini (3.8B)',            '~2.5GB RAM, ~2–3s response'),
        ('TTS engine',      'Piper TTS',                            'Offline, natural voice, <0.5s'),
        ('Database',        'SQLite + SQLAlchemy ORM',              'Tasks, meetings, projects, reminders'),
        ('Sheets sync',     'gspread Python library',               'OAuth2, auto-sync every 60s'),
        ('Scheduler',       'APScheduler',                          'Reminders, sync jobs, Pomodoro'),
        ('Process manager', 'systemd service',                      'Auto-start on boot, restart on crash'),
    ],
    col_widths=[1.5, 2.2, 2.5]
)

sub_heading(doc, '4.3  API Endpoints (FastAPI)')
make_table(doc,
    ['Method', 'Endpoint', 'Description'],
    [
        ('WS',    '/ws/audio',        'Receive PCM audio stream from ESP32, return WAV + JSON'),
        ('POST',  '/api/command',     'Text command fallback (testing / debug)'),
        ('GET',   '/api/tasks',       'Fetch task list (filter by sheet, status)'),
        ('GET',   '/api/projects',    'Fetch all sheets/projects'),
        ('POST',  '/api/sync',        'Force immediate Google Sheets sync'),
        ('GET',   '/api/status',      'Server health, model status, uptime'),
        ('GET',   '/api/reminders',   'Fetch pending reminders'),
        ('POST',  '/api/pomodoro',    'Start/stop/pause Pomodoro session'),
        ('GET',   '/api/wifi/scan',   'Scan WiFi networks (for setup)'),
        ('POST',  '/api/wifi/connect','Connect to new WiFi network'),
    ],
    col_widths=[0.7, 2.0, 3.5]
)

sub_heading(doc, '4.4  AI Pipeline Detail')
body(doc, 'Every voice command passes through this sequential pipeline on the Pi5:')
make_table(doc,
    ['Step', 'Component', 'Input', 'Output', 'Latency'],
    [
        ('1', 'faster-whisper',   'PCM audio (16kHz)',    'Transcript text',              '~1.5s'),
        ('2', 'Intent classifier','Transcript text',       'Intent + entity JSON',         '~2.0s'),
        ('3', 'Action executor',  'Intent JSON',           'DB write + Sheets sync',       '~0.3s'),
        ('4', 'Response gen',     'Action result',         'Natural language response',    '~1.5s'),
        ('5', 'Piper TTS',        'Response text',         'WAV audio file',               '~0.4s'),
        ('6', 'WebSocket send',   'WAV + JSON payload',    'Delivered to ESP32',           '~0.1s'),
        ('', 'TOTAL',             '',                      '',                             '~6s end-to-end'),
    ],
    col_widths=[0.5, 1.6, 1.8, 2.0, 1.0]
)

sub_heading(doc, '4.5  Ollama System Prompt Design')
body(doc, 'The LLM is constrained by a structured system prompt to prevent hallucination and ensure consistent JSON output:')
code_block(doc, '''You are VoiceDesk, a voice productivity assistant.
Extract structured data from user speech and return ONLY valid JSON.

Intents: ADD_TASK | DELETE_TASK | QUERY_TASKS | SCHEDULE_MEETING |
         NEW_SHEET | SWITCH_SHEET | SET_REMINDER | POMODORO_START |
         POMODORO_STOP | MORNING_BRIEFING | UNKNOWN

Output format:
{
  "intent": "ADD_TASK",
  "confidence": 0.95,
  "entities": {
    "title": "Review contract",
    "sheet": "Work",
    "due_date": "2026-06-05",
    "remarks": "Check clause 3 and 7"
  },
  "missing_fields": ["due_date"],
  "clarification_needed": "Which sheet should I add this to?",
  "suggested_followups": ["Set a reminder?", "Add renewal date?"]
}

Today's date: {current_date}. User's sheets: {sheet_list}.
NEVER guess. If a required field is missing, ask ONE clarifying question.''')

# ═══════════════════════════════════════════════════════════════════════════════
# 5. DATA ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '5', 'Data Architecture')

sub_heading(doc, '5.1  Database Schema (SQLite on Pi5)')
code_block(doc, '''-- Sheets = project buckets (each maps to a Google Sheets tab)
CREATE TABLE sheets (
  id            INTEGER PRIMARY KEY,
  name          TEXT NOT NULL,          -- "Work", "Personal", "Seed Form"
  google_sheet_id TEXT,                 -- Google Sheets file ID
  tab_name      TEXT,                   -- specific tab within the file
  color_tag     TEXT,                   -- UI colour for this sheet
  created_at    DATETIME DEFAULT CURRENT_TIMESTAMP,
  last_synced   DATETIME
);

-- Tasks
CREATE TABLE tasks (
  id            INTEGER PRIMARY KEY,
  sheet_id      INTEGER REFERENCES sheets(id),
  title         TEXT NOT NULL,
  status        TEXT DEFAULT 'pending', -- pending|in_progress|done|cancelled
  due_date      DATETIME,
  priority      TEXT DEFAULT 'normal',  -- low|normal|high|urgent
  remarks       TEXT,
  created_at    DATETIME DEFAULT CURRENT_TIMESTAMP,
  updated_at    DATETIME,
  synced        BOOLEAN DEFAULT 0
);

-- Meetings
CREATE TABLE meetings (
  id            INTEGER PRIMARY KEY,
  sheet_id      INTEGER REFERENCES sheets(id),
  title         TEXT NOT NULL,
  meeting_date  DATETIME,
  participants  TEXT,                   -- JSON array ["Rahul","John"]
  location      TEXT,
  remarks       TEXT,
  renewal_date  DATETIME,
  renewal_note  TEXT,
  created_at    DATETIME DEFAULT CURRENT_TIMESTAMP,
  synced        BOOLEAN DEFAULT 0
);

-- Reminders
CREATE TABLE reminders (
  id            INTEGER PRIMARY KEY,
  task_id       INTEGER,
  meeting_id    INTEGER,
  remind_at     DATETIME NOT NULL,
  message       TEXT,
  delivered     BOOLEAN DEFAULT 0
);

-- Command audit log
CREATE TABLE command_log (
  id            INTEGER PRIMARY KEY,
  transcript    TEXT,
  intent        TEXT,
  entities_json TEXT,
  success       BOOLEAN,
  response_text TEXT,
  latency_ms    INTEGER,
  created_at    DATETIME DEFAULT CURRENT_TIMESTAMP
);''')

sub_heading(doc, '5.2  Google Sheets Sync Strategy')
make_table(doc,
    ['Rule', 'Detail'],
    [
        ('Source of truth',     'SQLite on Pi5 — always primary'),
        ('Sync direction',      'Pi5 → Google Sheets (one-way push)'),
        ('Sync trigger',        'Every 60 seconds + immediately on every write'),
        ('Offline behaviour',   'Writes queue in SQLite; synced when internet restores'),
        ('Sheet mapping',       'One "sheet" in app = one tab in Google Spreadsheet'),
        ('Column structure',    'Title | Status | Due Date | Priority | Remarks | Created | Updated'),
        ('New sheet creation',  'Voice command creates new tab automatically'),
        ('Auth',                'OAuth2 service account — set up once, persists'),
    ],
    col_widths=[2.0, 4.2]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 6. NETWORK & COMMUNICATION
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '6', 'Network & Communication')

sub_heading(doc, '6.1  WebSocket Protocol')
body(doc, 'ESP32 maintains a persistent WebSocket connection to the Pi5 server. All communication is binary (audio) or JSON (commands/responses).')
code_block(doc, '''// ESP32 → Pi5: Audio packet
{
  "type": "audio_chunk",
  "session_id": "abc123",
  "sample_rate": 16000,
  "encoding": "PCM16",
  "data": "<base64 encoded audio>"
}

// ESP32 → Pi5: Button command
{
  "type": "command",
  "action": "BUTTON_OK",
  "context": "task_detail",
  "item_id": 42
}

// Pi5 → ESP32: Response
{
  "type": "response",
  "display": {
    "screen": "face",
    "state": "speaking",
    "text": "Added! Meeting with Rahul on Monday 3pm."
  },
  "audio_url": "/audio/response_abc123.wav",
  "follow_up": "Should I set a reminder before the meeting?"
}''')

sub_heading(doc, '6.2  Device Discovery')
make_table(doc,
    ['Method', 'Detail'],
    [
        ('mDNS hostname',    'Pi5 advertises as voicedesk.local — no IP configuration needed'),
        ('Fallback',         'Static IP configurable via NVS on ESP32 first-boot screen'),
        ('Port',             'FastAPI server on port 8000 (HTTP + WebSocket)'),
        ('Reconnect logic',  'ESP32 auto-reconnects with exponential backoff on disconnect'),
    ],
    col_widths=[1.8, 4.4]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 7. UI / UX DESIGN
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '7', 'UI / UX Design')

sub_heading(doc, '7.1  Screen Inventory')
make_table(doc,
    ['Screen', 'Trigger', 'Key Elements'],
    [
        ('Home / Idle Face',    'Default / inactivity timeout',         'Animated face (5 states), time, WiFi/sync status'),
        ('Main Menu',           'Short press BACK from home',           'Icons + labels: Tasks, Meetings, Projects, Pomodoro, Settings, WiFi'),
        ('Task List',           'Select "Tasks" in menu',               'Scrollable list, status icons, due dates'),
        ('Task Detail',         'OK on a task',                         'Full info + action options (Done / Delete / Edit remarks)'),
        ('Listening',           'MIC button pressed / wake word',       'Wide-eye face, audio level bar, "Listening..." text'),
        ('Thinking',            'Audio received, processing',           'Squinting face, animated dots'),
        ('Speaking / Response', 'TTS audio playing',                    'Happy face, response text scroll, mouth animation'),
        ('Pomodoro Timer',      'Select Pomodoro in menu / voice cmd',  'Countdown, progress bar, session counter, task name'),
        ('WiFi Setup',          'First boot / Settings > WiFi',         'Network scan list, signal strength, voice password entry'),
        ('Error',               'Server unreachable / parse failure',   'Sad face, error message, retry prompt'),
    ],
    col_widths=[1.8, 2.0, 2.4]
)

sub_heading(doc, '7.2  Animated Face States')
make_table(doc,
    ['State', 'Eyes', 'Mouth', 'When'],
    [
        ('IDLE',      '◕  ◕  (blinking every 3–5s)', '‿  slight smile',          'Waiting for input'),
        ('LISTENING', '◉  ◉  (wide, no blink)',       '○  open',                  'MIC button held / wake word active'),
        ('THINKING',  '◑  ◑  (slow squint)',          '...  dots animate',        'Audio processing on Pi5'),
        ('SPEAKING',  '◕  ◕  (happy)',                '◡◡◡  bounces with audio',  'TTS audio playing back'),
        ('SUCCESS',   '◕  ◉  (wink)',                 '◡  big smile',             'Task added / action confirmed'),
        ('ERROR',     '◔  ◔  (worried)',              '︵  frown',                'Error / unreachable server'),
    ],
    col_widths=[1.2, 1.8, 2.0, 1.8]
)

note_box(doc, 'Face animations are rendered as TFT_eSPI sprites — pre-drawn into a 100×100px off-screen buffer and blitted to display each frame. Target: 10–15fps for smooth blink/mouth animation.')

# ═══════════════════════════════════════════════════════════════════════════════
# 8. VOICE COMMAND FLOW
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '8', 'Voice Command Flow Examples')

sub_heading(doc, '8.1  Schedule a Meeting')
code_block(doc, '''User:    "Schedule a meeting with Rahul next Monday at 3pm"
STT:      "Schedule a meeting with Rahul next Monday at 3pm"
Intent:   SCHEDULE_MEETING
Entities: { title: "Meeting with Rahul", date: "2026-06-08", time: "15:00" }
Missing:  sheet/project
Response: "Which project should I add this to — Work, Personal, or Seed Form?"
User:     "Work"
Action:   INSERT meetings (sheet=Work, date=2026-06-08 15:00)
          SYNC → Google Sheets tab "Work"
Follow-up: "Done! Should I set a reminder before the meeting?
            Also, want to add a renewal or follow-up date?"''')

sub_heading(doc, '8.2  Add a Task')
code_block(doc, '''User:    "Add task review seed form documents by Friday"
Intent:   ADD_TASK
Entities: { title: "Review Seed Form documents", due: "2026-06-05",
            sheet_hint: "Seed Form" }
Action:   INSERT tasks (sheet="Seed Form", due=Friday)
Response: "Added to Seed Form sheet, due this Friday. Any remarks to add?"''')

sub_heading(doc, '8.3  Query Tasks')
code_block(doc, '''User:    "What's pending in Work?"
Intent:   QUERY_TASKS
Entities: { sheet: "Work", status: "pending" }
Action:   SELECT tasks WHERE sheet=Work AND status=pending
Response: "You have 3 pending tasks in Work:
           1. Send proposal — due tomorrow
           2. Review contract — no due date
           3. Follow up with client — overdue 2 days"
Display:  Task list rendered on ESP32 screen''')

sub_heading(doc, '8.4  Start Pomodoro')
code_block(doc, '''User:    "Start a Pomodoro for review contract"
Intent:   POMODORO_START
Entities: { task: "Review contract", duration: 25 }
Action:   Set timer, update display to Pomodoro screen
Response: "Pomodoro started! 25 minutes for Review contract.
           I'll let you know when it's time for a break."''')

# ═══════════════════════════════════════════════════════════════════════════════
# 9. GOOGLE SHEETS INTEGRATION
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '9', 'Google Sheets Integration')

sub_heading(doc, '9.1  Setup (One-Time)')
make_table(doc,
    ['Step', 'Action'],
    [
        ('1', 'Create a Google Cloud project and enable Google Sheets API'),
        ('2', 'Create a Service Account, download credentials.json'),
        ('3', 'Copy credentials.json to Pi5 at /home/pi/voicedesk/credentials.json'),
        ('4', 'Share your Google Spreadsheet with the service account email'),
        ('5', 'Run: python setup_sheets.py — confirms connection and creates default tabs'),
        ('6', 'Done — all future access is automatic'),
    ],
    col_widths=[0.5, 5.7]
)

sub_heading(doc, '9.2  Sheet Structure per Project')
code_block(doc, '''Google Spreadsheet: "VoiceDesk Data"
  ├── Tab: Work          ← tasks + meetings for Work
  ├── Tab: Personal      ← personal tasks
  ├── Tab: Seed Form     ← Seed Form project
  └── Tab: [New Sheet]   ← created by voice command

Each tab columns:
  A: ID | B: Type | C: Title | D: Status | E: Due Date |
  F: Priority | G: Remarks | H: Created | I: Updated | J: Synced At''')

# ═══════════════════════════════════════════════════════════════════════════════
# 10. RASPBERRY PI 5 SERVER SETUP
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '10', 'Raspberry Pi 5 Server Setup')

sub_heading(doc, '10.1  Required Upgrades')
make_table(doc,
    ['Upgrade', 'Component', 'Cost', 'Priority', 'Reason'],
    [
        ('NVMe SSD',     'Official Pi5 M.2 HAT+ + 256GB NVMe M.2 2242', '~$35', '🔴 Critical', 'SD card too slow for Whisper + DB writes. 10× faster.'),
        ('Active cooler','Official Pi5 Active Cooler (fan + heatsink)',   '~$8',  '🔴 Critical', 'Pi5 throttles under sustained AI load without cooling.'),
        ('Power supply', '27W USB-C PD official supply',                  '~$12', '🔴 Critical', 'Stable power for NVMe + heavy CPU workloads.'),
        ('UPS HAT',      'PiJuice / Waveshare UPS HAT',                   '~$25', '🟡 Optional', 'Keeps server alive during power cuts.'),
    ],
    col_widths=[1.2, 2.2, 0.7, 0.9, 2.2]
)

sub_heading(doc, '10.2  Ollama Model Comparison')
make_table(doc,
    ['Model', 'Params', 'RAM', 'Speed (Pi5)', 'Recommendation'],
    [
        ('phi3:mini',    '3.8B', '~2.5GB', '2–3 sec',  '✅ Best choice — fast, smart enough for intent'),
        ('llama3.2:3b',  '3B',   '~2.0GB', '1–2 sec',  '⚡ Fastest option, slightly less accurate'),
        ('mistral:7b',   '7B',   '~4.5GB', '6–8 sec',  '⚠️ Too slow for Pi5 4GB; needs 8GB model'),
        ('gemma2:2b',    '2B',   '~1.8GB', '1 sec',    '🔄 Backup if phi3 RAM is too tight'),
    ],
    col_widths=[1.4, 0.8, 0.8, 1.3, 2.9]
)

sub_heading(doc, '10.3  Installation Commands')
code_block(doc, '''# 1. Install OS (Raspberry Pi OS Lite 64-bit) on NVMe via rpi-imager

# 2. Enable SSH, connect to network, then:
sudo apt update && sudo apt upgrade -y

# 3. Install Python dependencies
pip install fastapi uvicorn faster-whisper piper-tts gspread \
            SQLAlchemy APScheduler sounddevice websockets

# 4. Install Ollama
curl -fsSL https://ollama.ai/install.sh | sh
ollama pull phi3:mini

# 5. Set up mDNS hostname
sudo hostnamectl set-hostname voicedesk
# voicedesk.local now resolves on LAN

# 6. Install as systemd service
sudo systemctl enable voicedesk
sudo systemctl start voicedesk

# 7. Google Sheets setup
python setup_sheets.py --credentials credentials.json''')

# ═══════════════════════════════════════════════════════════════════════════════
# 11. COMPLETE WIRING REFERENCE
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '11', 'Complete Wiring Reference')

sub_heading(doc, '11.1  MPI3501 Display → ESP32')
make_table(doc,
    ['Display Pin', 'ESP32 GPIO', 'Physical Pin', 'Notes'],
    [
        ('+5V',   'VIN',     '—',      'From 5V power rail'),
        ('GND',   'GND',     '—',      ''),
        ('DC',    'GPIO 2',  'Pin 26', 'Data/Command select'),
        ('RST',   'GPIO 4',  'Pin 26', 'Display reset'),
        ('CS',    'GPIO 15', 'Pin 13', 'Display chip select'),
        ('MOSI',  'GPIO 23', 'Pin 37', 'SPI data out'),
        ('SCK',   'GPIO 18', 'Pin 30', 'SPI clock'),
        ('T_CS',  '3.3V',   '—',      'Tie HIGH — disables touch'),
        ('MISO',  'NC',      '—',      'Not needed (no touch)'),
        ('IRQ',   'NC',      '—',      'Not needed (no touch)'),
    ],
    col_widths=[1.2, 1.2, 1.2, 2.6]
)

sub_heading(doc, '11.2  INMP441 Microphone → ESP32')
make_table(doc,
    ['INMP441 Pin', 'ESP32 GPIO', 'Notes'],
    [
        ('VDD',  '3.3V',    '⚠️ 3.3V ONLY — not 5V tolerant'),
        ('GND',  'GND',     ''),
        ('SD',   'GPIO 32', 'I2S0 data input'),
        ('WS',   'GPIO 25', 'I2S0 LRCK / word select'),
        ('SCK',  'GPIO 26', 'I2S0 BCLK / bit clock'),
        ('L/R',  'GND',     'GND = left channel; 3.3V = right channel'),
    ],
    col_widths=[1.4, 1.4, 3.4]
)

sub_heading(doc, '11.3  MAX98357A Speaker → ESP32')
make_table(doc,
    ['MAX98357A Pin', 'ESP32 GPIO', 'Notes'],
    [
        ('VIN',   '5V',     'From 5V power rail — louder output'),
        ('GND',   'GND',    ''),
        ('BCLK',  'GPIO 27','I2S1 bit clock'),
        ('LRC',   'GPIO 25','I2S1 LRCK — can share with INMP441 WS'),
        ('DIN',   'GPIO 22','I2S1 data output'),
        ('SD',    'Float',  'Leave floating = always on; GND = mute'),
        ('GAIN',  'Float',  'Float=9dB, GND=12dB, VIN=15dB'),
        ('Speaker+/−', '8Ω 3W speaker', 'Direct connect, no extra components'),
    ],
    col_widths=[1.6, 1.4, 3.2]
)

sub_heading(doc, '11.4  Buttons → ESP32')
make_table(doc,
    ['Button', 'GPIO', 'Wiring'],
    [
        ('BACK',    'GPIO 34', 'One side to GPIO 34, other to GND. INPUT_PULLUP.'),
        ('UP',      'GPIO 35', 'One side to GPIO 35, other to GND. INPUT_PULLUP.'),
        ('DOWN',    'GPIO 33', 'One side to GPIO 33, other to GND. INPUT_PULLUP.'),
        ('OK',      'GPIO 13', 'One side to GPIO 13, other to GND. INPUT_PULLUP.'),
        ('MIC/PTT', 'GPIO 14', 'One side to GPIO 14, other to GND. INPUT_PULLUP.'),
    ],
    col_widths=[1.0, 1.0, 4.2]
)

sub_heading(doc, '11.5  Complete ESP32 GPIO Summary')
make_table(doc,
    ['GPIO', 'Function', 'Direction'],
    [
        ('GPIO 2',  'Display DC',          'OUT'),
        ('GPIO 4',  'Display RST',         'OUT'),
        ('GPIO 13', 'BTN OK',              'IN (PULLUP)'),
        ('GPIO 14', 'BTN MIC/PTT',         'IN (PULLUP)'),
        ('GPIO 15', 'Display CS',          'OUT'),
        ('GPIO 18', 'SPI CLK',             'OUT'),
        ('GPIO 19', 'SPI MISO (unused)',   'IN'),
        ('GPIO 22', 'MAX98357A DIN',       'OUT (I2S1)'),
        ('GPIO 23', 'SPI MOSI',            'OUT'),
        ('GPIO 25', 'I2S LRCK (shared)',   'OUT'),
        ('GPIO 26', 'INMP441 BCLK',        'OUT (I2S0)'),
        ('GPIO 27', 'MAX98357A BCLK',      'OUT (I2S1)'),
        ('GPIO 32', 'INMP441 DATA',        'IN  (I2S0)'),
        ('GPIO 33', 'BTN DOWN',            'IN (PULLUP)'),
        ('GPIO 34', 'BTN BACK',            'IN only (PULLUP)'),
        ('GPIO 35', 'BTN UP',              'IN only (PULLUP)'),
    ],
    col_widths=[1.0, 2.5, 1.8]
)
note_box(doc, 'GPIO 34 and 35 are input-only on ESP32 — ideal for buttons. No pin conflicts exist across all three peripherals.')

# ═══════════════════════════════════════════════════════════════════════════════
# 12. BILL OF MATERIALS
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '12', 'Bill of Materials')

sub_heading(doc, '12.1  ESP32 Handheld Device')
make_table(doc,
    ['#', 'Component', 'Model / Spec', 'Qty', 'Est. Cost'],
    [
        ('1',  'Microcontroller',     'ESP32 DevKit v1 (ESP-WROOM-32, 38-pin)',          '1', '$5–8'),
        ('2',  'Display',             'MPI3501 3.5" ILI9486 SPI TFT 480×320 (owned)',   '1', '$0 (owned)'),
        ('3',  'Microphone',          'INMP441 I2S MEMS (owned)',                        '1', '$0 (owned)'),
        ('4',  'Speaker amplifier',   'MAX98357A I2S 3W breakout (owned)',               '1', '$0 (owned)'),
        ('5',  'Speaker',             '40mm 8Ω 3W',                                     '1', '$1–3'),
        ('6',  'Tactile buttons',     '12×12mm momentary push button',                  '5', '$0.50'),
        ('7',  'Battery',             '3.7V Li-Po 3000mAh with protection',             '1', '$5–8'),
        ('8',  'Charge module',       'TP4056 USB-C with protection',                   '1', '$0.80'),
        ('9',  'Boost converter',     'MT3608 3.7V→5V 2A (or IP5306 all-in-one)',       '1', '$1–2'),
        ('10', 'Power switch',        'SPDT slide switch',                              '1', '$0.50'),
        ('11', 'LED indicators',      '3mm LED green + blue + 220Ω resistors',         '2', '$0.30'),
        ('12', 'Capacitors',          '100µF + 10µF electrolytic',                      '2', '$0.30'),
        ('13', 'Perfboard / PCB',     '7×5cm perfboard',                               '1', '$0.50'),
        ('14', 'Dupont wires',        'Female-female 20cm jumpers (pack)',               '1', '$1.00'),
        ('15', 'Enclosure',           '3D printed or 100×65×30mm project box',          '1', '$2–5'),
        ('',   'SUBTOTAL',            '',                                               '',  '~$18–30'),
    ],
    col_widths=[0.4, 1.8, 2.5, 0.4, 0.9]
)

sub_heading(doc, '12.2  Raspberry Pi 5 Upgrades')
make_table(doc,
    ['#', 'Component', 'Spec', 'Qty', 'Est. Cost'],
    [
        ('1', 'M.2 HAT+',      'Official Raspberry Pi M.2 HAT+',                 '1', '$12–15'),
        ('2', 'NVMe SSD',      '256GB M.2 2242 NVMe (WD Green / Kingston)',      '1', '$20–25'),
        ('3', 'Active cooler', 'Official Pi5 Active Cooler',                     '1', '$5–8'),
        ('4', 'Power supply',  '27W USB-C PD official supply',                   '1', '$12'),
        ('5', 'Ethernet cable','Cat5e 1m',                                       '1', '$2'),
        ('',  'SUBTOTAL',      '',                                               '',  '~$51–62'),
    ],
    col_widths=[0.4, 1.8, 2.8, 0.4, 0.9]
)

sub_heading(doc, '12.3  Optional Enhancements')
make_table(doc,
    ['Component', 'Use', 'Cost'],
    [
        ('Vibration motor (ERM)',   'Haptic feedback on button press',              '$1'),
        ('WS2812 RGB LED',          'Status ring — colour shows device state',      '$1'),
        ('DS3231 RTC module',       'Keep time without WiFi',                       '$2'),
        ('Passive buzzer',          'Beep on wake word / Pomodoro bell',            '$0.50'),
        ('PiSugar S (Pi Zero)',     'If Pi Zero used as client instead of ESP32',  '$18–25'),
        ('UPS HAT for Pi5',         'Power backup for server',                     '$20–25'),
    ],
    col_widths=[2.2, 2.8, 1.2]
)

body(doc, 'Grand Total (excluding items already owned): approximately $69–92 USD')

# ═══════════════════════════════════════════════════════════════════════════════
# 13. DEVELOPMENT ROADMAP
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '13', 'Development Roadmap')

make_table(doc,
    ['Phase', 'Feature', 'Priority', 'Est. Effort'],
    [
        ('Phase 1 — MVP',      'ESP32 firmware + display + buttons',                  '🔴 P0', '1 week'),
        ('Phase 1 — MVP',      'Pi5 FastAPI server + WebSocket audio',                '🔴 P0', '1 week'),
        ('Phase 1 — MVP',      'Whisper STT integration',                             '🔴 P0', '2 days'),
        ('Phase 1 — MVP',      'Ollama phi3 intent parser',                           '🔴 P0', '3 days'),
        ('Phase 1 — MVP',      'SQLite DB + basic task CRUD',                         '🔴 P0', '2 days'),
        ('Phase 1 — MVP',      'Google Sheets sync (one sheet)',                      '🔴 P0', '2 days'),
        ('Phase 1 — MVP',      'Piper TTS + playback on ESP32',                       '🔴 P0', '2 days'),
        ('Phase 1 — MVP',      'Animated face UI (5 states)',                         '🔴 P0', '3 days'),
        ('Phase 2 — Smart',    'Multi-sheet support (create/switch by voice)',         '🟡 P1', '3 days'),
        ('Phase 2 — Smart',    'Meeting scheduling + reminders',                      '🟡 P1', '3 days'),
        ('Phase 2 — Smart',    'Renewal / follow-up smart suggestions',               '🟡 P1', '2 days'),
        ('Phase 2 — Smart',    'Pomodoro timer (voice start/stop)',                   '🟡 P1', '2 days'),
        ('Phase 2 — Smart',    'Morning briefing ("What is my day?")',                '🟡 P1', '2 days'),
        ('Phase 2 — Smart',    'WiFi setup via voice (password spoken)',              '🟡 P1', '2 days'),
        ('Phase 3 — Advanced', 'Recurring tasks + task templates',                   '🟢 P2', '3 days'),
        ('Phase 3 — Advanced', 'Google Calendar integration',                        '🟢 P2', '4 days'),
        ('Phase 3 — Advanced', 'Web dashboard (view from phone browser)',            '🟢 P2', '1 week'),
        ('Phase 3 — Advanced', 'Offline mode (full local fallback)',                 '🟢 P2', '3 days'),
        ('Phase 3 — Advanced', 'OTA firmware update for ESP32',                      '🟢 P2', '2 days'),
    ],
    col_widths=[1.5, 3.0, 0.8, 0.9]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 14. ARCHITECTURAL DECISIONS
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '14', 'Architectural Decisions & Trade-offs')

make_table(doc,
    ['Decision', 'Choice Made', 'Alternatives Considered', 'Rationale'],
    [
        ('Handheld MCU',
         'ESP32',
         'Pi Zero 2W, Pi Pico W',
         'Instant boot, low power, dual I2S buses, WiFi built-in. Pi Zero overkill for thin client.'),
        ('AI server',
         'Raspberry Pi 5',
         'Cloud APIs, Pi 4',
         'Full local privacy. Pi5 4× faster than Pi4 for inference. Cloud has latency + cost.'),
        ('LLM model',
         'phi3:mini via Ollama',
         'mistral:7b, gemma2:2b',
         'Best speed/accuracy balance for 4GB Pi5. Mistral too slow; gemma2 less accurate.'),
        ('STT engine',
         'faster-whisper base.en',
         'Vosk, Google STT, Whisper large',
         'faster-whisper is 4× faster than original Whisper with same accuracy. Fully local.'),
        ('TTS engine',
         'Piper TTS',
         'espeak, Google TTS, Coqui',
         'Most natural voice of local options. <500ms latency. Fully offline.'),
        ('Display control',
         '5 physical buttons',
         'Resistive touch (XPT2046), capacitive touch',
         'Buttons work with gloves, in dark, without looking at screen. Voice is primary anyway.'),
        ('Data layer',
         'SQLite + Google Sheets sync',
         'PostgreSQL, Firebase, Notion',
         'SQLite = zero setup, offline-first. Google Sheets = user already familiar, shareable.'),
        ('Communication',
         'WebSocket (persistent)',
         'HTTP polling, MQTT',
         'WebSocket enables real-time audio streaming and instant push responses from Pi5.'),
        ('Wake word',
         'Porcupine on-device',
         'Always-stream to Pi5, button-only',
         'On-device wake word = privacy (no audio sent until triggered). PTT as backup.'),
    ],
    col_widths=[1.3, 1.2, 1.5, 2.2]
)

para_space(doc, after=20)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— End of Document —')
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = C_GRAY

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/sessions/sharp-affectionate-dijkstra/mnt/outputs/VoiceDesk_Architecture.docx'
doc.save(out)
print(f'Saved: {out}')
