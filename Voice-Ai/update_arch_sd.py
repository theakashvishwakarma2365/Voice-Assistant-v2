"""
Update VoiceDesk_Architecture.docx with SD card offline storage section.
Changes:
  1. TOC  — add entry "15  SD Card Offline Storage & Sync"
  2. Section 2.2 data flow — add SD card branch
  3. Section 3.1 component table — add SD card row
  4. Section 4.1 firmware stack — add SD.h row
  5. Section 11.5 GPIO summary — add GPIO 5 row
  6. Section 12.1 BOM — add SD module row + update subtotal
  7. NEW Section 15 — full SD card section (wiring, logic, code, indicators)
  8. Section 14 decisions table — add SD card decision row
  9. Update meta status line
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import re

doc = Document('/sessions/sharp-affectionate-dijkstra/mnt/outputs/VoiceDesk_Architecture.docx')

# ── Colour palette (must match original) ─────────────────────────────────────
C_DARK    = RGBColor(0x1A, 0x1A, 0x2E)
C_ACCENT  = RGBColor(0x0F, 0x3C, 0x78)
C_LIGHT   = RGBColor(0xE8, 0xF0, 0xFE)
C_GREEN   = RGBColor(0x1B, 0x87, 0x54)
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_GRAY    = RGBColor(0x64, 0x74, 0x87)
C_TABLE_H = RGBColor(0x0F, 0x3C, 0x78)
C_TABLE_A = RGBColor(0xF0, 0xF4, 0xFF)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)

def new_para(text='', size=10, bold=False, italic=False,
             color=None, align=None, indent_cm=None,
             space_before=2, space_after=4, courier=False):
    from docx.oxml import OxmlElement as OE
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    if align:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), align)
        pPr.append(jc)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(space_before * 20)))
    spacing.set(qn('w:after'),  str(int(space_after  * 20)))
    pPr.append(spacing)
    if indent_cm:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(int(indent_cm * 720)))
        pPr.append(ind)
    p.append(pPr)
    if text:
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(size * 2)))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(size * 2)))
        rPr.append(szCs)
        if bold:
            rPr.append(OxmlElement('w:b'))
        if italic:
            rPr.append(OxmlElement('w:i'))
        if courier:
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'),    'Courier New')
            rFonts.set(qn('w:hAnsi'),    'Courier New')
            rFonts.set(qn('w:cs'),       'Courier New')
            rPr.append(rFonts)
        if color:
            clr = OxmlElement('w:color')
            clr.set(qn('w:val'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
            rPr.append(clr)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
    return p

def make_heading_para(number, title):
    """Section heading matching original style."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(14 * 20))
    spacing.set(qn('w:after'),  str(4  * 20))
    pPr.append(spacing)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:color'), '0F3C78')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    b  = OxmlElement('w:b');  rPr.append(b)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '28'); rPr.append(sz)
    clr = OxmlElement('w:color')
    clr.set(qn('w:val'), '0F3C78')
    rPr.append(clr)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = f'{number}  {title}'
    r.append(t)
    p.append(r)
    return p

def make_sub_heading_para(title):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(10 * 20))
    spacing.set(qn('w:after'),  str(3  * 20))
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    b  = OxmlElement('w:b');  rPr.append(b)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '22'); rPr.append(sz)
    clr = OxmlElement('w:color'); clr.set(qn('w:val'), '1A1A2E'); rPr.append(clr)
    r.append(rPr)
    t = OxmlElement('w:t'); t.text = title; r.append(t)
    p.append(r)
    return p

def make_note_para(text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(4 * 20))
    spacing.set(qn('w:after'),  str(6 * 20))
    pPr.append(spacing)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(int(0.5 * 720)))
    ind.set(qn('w:right'), str(int(0.5 * 720)))
    pPr.append(ind)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'E8F0FE')
    pPr.append(shd)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    i   = OxmlElement('w:i');  rPr.append(i)
    sz  = OxmlElement('w:sz'); sz.set(qn('w:val'), '19'); rPr.append(sz)
    clr = OxmlElement('w:color'); clr.set(qn('w:val'), '0F3C78'); rPr.append(clr)
    r.append(rPr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve')
    t.text = f'  ℹ  {text}'; r.append(t)
    p.append(r)
    return p

def code_paras(code_text):
    paras = []
    for line in code_text.strip('\n').split('\n'):
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'),  '0')
        pPr.append(spacing)
        ind = OxmlElement('w:ind'); ind.set(qn('w:left'), str(int(0.8*720))); pPr.append(ind)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),  'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F3F4F6'); pPr.append(shd)
        p.append(pPr)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Courier New'); rFonts.set(qn('w:hAnsi'), 'Courier New')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '17'); rPr.append(sz)
        clr = OxmlElement('w:color'); clr.set(qn('w:val'), '1E1E1E'); rPr.append(clr)
        r.append(rPr)
        t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve')
        t.text = line if line else ' '; r.append(t)
        p.append(r)
        paras.append(p)
    # trailing spacer
    sp = OxmlElement('w:p')
    spPr = OxmlElement('w:pPr')
    spacing = OxmlElement('w:spacing'); spacing.set(qn('w:before'), '0'); spacing.set(qn('w:after'), str(4*20)); spPr.append(spacing)
    sp.append(spPr); paras.append(sp)
    return paras

def make_table_xml(doc_obj, headers, rows, col_widths=None):
    """Build a proper docx table and return the tbl XML element."""
    from docx.shared import Inches
    tbl_obj = doc_obj.add_table(rows=1+len(rows), cols=len(headers))
    tbl_obj.style = 'Table Grid'
    tbl_obj.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl_obj.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, C_TABLE_H)
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True; run.font.color.rgb = C_WHITE; run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        bg = C_TABLE_A if ri % 2 == 0 else C_WHITE
        for ci, val in enumerate(row):
            cell = tbl_obj.rows[ri+1].cells[ci]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9); run.font.color.rgb = C_DARK
    if col_widths:
        for row in tbl_obj.rows:
            for ci, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[ci])
    # detach from doc body and return raw XML element
    tbl_elem = tbl_obj._tbl
    tbl_elem.getparent().remove(tbl_elem)
    return tbl_elem

def insert_after(body, anchor_elem, *new_elems):
    """Insert new_elems immediately after anchor_elem in body."""
    idx = list(body).index(anchor_elem)
    for offset, elem in enumerate(new_elems):
        body.insert(idx + 1 + offset, elem)

def find_para(doc_obj, text_contains):
    for p in doc_obj.paragraphs:
        if text_contains in p.text:
            return p
    return None

def find_table_after_para(doc_obj, para):
    """Return the first table element that appears after para in the body."""
    body = doc_obj.element.body
    elems = list(body)
    try:
        start = elems.index(para._p)
    except ValueError:
        return None
    for elem in elems[start+1:]:
        if elem.tag.endswith('}tbl'):
            # wrap back to a docx Table object
            from docx.table import Table
            return Table(elem, doc_obj)
    return None

# ══════════════════════════════════════════════════════════════════════════════
# 1. UPDATE TOC — add "15  SD Card Offline Storage & Sync"
# ══════════════════════════════════════════════════════════════════════════════
toc_anchor = find_para(doc, '14  Architectural Decisions & Trade-offs')
if toc_anchor:
    new_toc = new_para('15  SD Card Offline Storage & Sync',
                       size=10, color=C_DARK, space_before=1, space_after=1,
                       indent_cm=0.5)
    insert_after(doc.element.body, toc_anchor._p, new_toc)
    print('✓ TOC entry added')

# ══════════════════════════════════════════════════════════════════════════════
# 2. UPDATE SECTION 3.1 Component Table — add SD card row
# ══════════════════════════════════════════════════════════════════════════════
comp_heading = find_para(doc, '3.1  ESP32 Handheld Device — Component Roles')
if comp_heading:
    tbl = find_table_after_para(doc, comp_heading)
    if tbl:
        new_row = tbl.add_row()
        bg = C_TABLE_A  # even row
        vals = ('MicroSD Card Module', 'SPI MicroSD breakout w/ 3.3V reg', 'SPI', 'Offline task queue, auto/manual sync to Pi5')
        for ci, val in enumerate(vals):
            cell = new_row.cells[ci]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9); run.font.color.rgb = C_DARK
        print('✓ Component table row added')

# ══════════════════════════════════════════════════════════════════════════════
# 3. UPDATE SECTION 4.1 Firmware Stack — add SD.h row
# ══════════════════════════════════════════════════════════════════════════════
fw_heading = find_para(doc, '4.1  ESP32 Firmware Stack')
if fw_heading:
    tbl = find_table_after_para(doc, fw_heading)
    if tbl:
        new_row = tbl.add_row()
        vals = ('Offline storage', 'SD.h (Arduino built-in)',
                'MicroSD on shared SPI bus; JSON queue files; auto-sync on reconnect')
        for ci, val in enumerate(vals):
            cell = new_row.cells[ci]
            set_cell_bg(cell, C_WHITE)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9); run.font.color.rgb = C_DARK
        print('✓ Firmware stack table row added')

# ══════════════════════════════════════════════════════════════════════════════
# 4. UPDATE SECTION 11.5 GPIO Summary — add GPIO 5 row
# ══════════════════════════════════════════════════════════════════════════════
gpio_heading = find_para(doc, '11.5  Complete ESP32 GPIO Summary')
if gpio_heading:
    tbl = find_table_after_para(doc, gpio_heading)
    if tbl:
        # Insert GPIO 5 at top of data rows (after header)
        # We'll just add it as a new row — sorted appearance isn't critical
        new_row = tbl.add_row()
        vals = ('GPIO 5', 'SD Card CS', 'OUT')
        for ci, val in enumerate(vals):
            cell = new_row.cells[ci]
            set_cell_bg(cell, C_TABLE_A)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9); run.font.color.rgb = C_DARK
        print('✓ GPIO table row added')

# ══════════════════════════════════════════════════════════════════════════════
# 5. UPDATE SECTION 12.1 BOM — add SD module row + update subtotal text
# ══════════════════════════════════════════════════════════════════════════════
bom_heading = find_para(doc, '12.1  ESP32 Handheld Device')
if bom_heading:
    tbl = find_table_after_para(doc, bom_heading)
    if tbl:
        # Find the SUBTOTAL row and insert SD card before it
        subtotal_row_idx = None
        for ri, row in enumerate(tbl.rows):
            if 'SUBTOTAL' in row.cells[1].text:
                subtotal_row_idx = ri
                break
        if subtotal_row_idx is not None:
            # Add new row (appended, then we'll update subtotal text)
            new_row = tbl.add_row()
            vals = ('16', 'MicroSD card module', 'SPI breakout with 3.3V regulator + level shift', '1', '$1–2')
            for ci, val in enumerate(vals):
                cell = new_row.cells[ci]
                set_cell_bg(cell, C_TABLE_A)
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(9); run.font.color.rgb = C_DARK
            # Update subtotal row
            sub_cell = tbl.rows[subtotal_row_idx].cells[4]
            for p in sub_cell.paragraphs:
                for r in p.runs:
                    r.text = '~$19–32'
            print('✓ BOM table row added + subtotal updated')

# ══════════════════════════════════════════════════════════════════════════════
# 6. UPDATE Grand Total paragraph
# ══════════════════════════════════════════════════════════════════════════════
grand_total_p = find_para(doc, 'Grand Total (excluding items already owned)')
if grand_total_p:
    for run in grand_total_p.runs:
        run.text = run.text.replace('$69–92', '$70–94')
    print('✓ Grand total updated')

# ══════════════════════════════════════════════════════════════════════════════
# 7. UPDATE SECTION 14 decisions table — add SD card row
# ══════════════════════════════════════════════════════════════════════════════
dec_heading = find_para(doc, '14  Architectural Decisions & Trade-offs')
# Find the second one (TOC is first, actual section is second)
all_dec = [p for p in doc.paragraphs if '14  Architectural Decisions' in p.text]
if len(all_dec) >= 2:
    tbl = find_table_after_para(doc, all_dec[1])
    if tbl:
        new_row = tbl.add_row()
        bg = C_TABLE_A
        vals = (
            'Offline storage',
            'MicroSD + SD.h',
            'No offline mode, cloud-only',
            'Enables full operation without internet. Shared SPI bus keeps pin count low. Auto-sync on reconnect.'
        )
        for ci, val in enumerate(vals):
            cell = new_row.cells[ci]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9); run.font.color.rgb = C_DARK
        print('✓ Decision table row added')

# ══════════════════════════════════════════════════════════════════════════════
# 8. UPDATE Section 2.2 Data Flow — insert SD branch note
# ══════════════════════════════════════════════════════════════════════════════
# Find the note box after data flow and update its text
flow_note = find_para(doc, 'All processing is local. Internet is only needed')
if flow_note:
    for run in flow_note.runs:
        if 'All processing is local' in run.text:
            run.text = ('  ℹ  All processing is local. Internet is only needed for Google Sheets sync. '
                        'When WiFi/Pi5 is unreachable, all data is queued to the MicroSD card and '
                        'auto-synced when connection is restored. Manual sync is also available via '
                        'long-press OK or voice command "sync my data".')
    print('✓ Data flow note updated')

# ══════════════════════════════════════════════════════════════════════════════
# 9. INSERT NEW SECTION 15 before the "End of Document" paragraph
# ══════════════════════════════════════════════════════════════════════════════
end_para = find_para(doc, '— End of Document —')
body = doc.element.body

def insert_before(body, anchor_elem, *new_elems):
    idx = list(body).index(anchor_elem)
    for offset, elem in enumerate(new_elems):
        body.insert(idx + offset, elem)

# Build all new elements for Section 15
new_elems = []

# Page break
pb = OxmlElement('w:p')
pbPr = OxmlElement('w:pPr'); pb.append(pbPr)
pbR  = OxmlElement('w:r')
pbRPr = OxmlElement('w:rPr'); pbR.append(pbRPr)
pbBr = OxmlElement('w:br')
pbBr.set(qn('w:type'), 'page')
pbR.append(pbBr); pb.append(pbR)
new_elems.append(pb)

# Section heading
new_elems.append(make_heading_para('15', 'SD Card Offline Storage & Sync'))

# Intro body
new_elems.append(new_para(
    'The ESP32 includes a MicroSD card module on the shared SPI bus. When the device cannot '
    'reach the Pi5 server (WiFi down, server rebooting, out of range), all voice commands are '
    'captured, processed locally where possible, and queued as JSON files on the SD card. '
    'When connectivity is restored, the queue is automatically flushed to the Pi5 and synced '
    'to Google Sheets.',
    size=10, color=C_DARK
))

# 15.1 Wiring
new_elems.append(make_sub_heading_para('15.1  MicroSD Module Wiring → ESP32'))

# Build wiring table in a temp doc then steal the element
tmp = Document()
wiring_tbl = make_table_xml(tmp,
    ['SD Module Pin', 'ESP32 GPIO', 'Notes'],
    [
        ('VCC',   '3.3V',    '⚠️ Use module with onboard 3.3V regulator + level shifter'),
        ('GND',   'GND',     ''),
        ('CS',    'GPIO 5',  'Dedicated SD chip select — new pin'),
        ('MOSI',  'GPIO 23', 'Shared with display SPI bus'),
        ('SCK',   'GPIO 18', 'Shared with display SPI bus'),
        ('MISO',  'GPIO 19', 'Shared with display SPI bus'),
    ],
    col_widths=[1.4, 1.2, 3.6]
)
new_elems.append(wiring_tbl)
new_elems.append(new_para('', size=10, space_before=0, space_after=4))

new_elems.append(make_note_para(
    'GPIO 5 is the only new pin needed. MOSI/MISO/SCK are already wired for the display. '
    'Display CS (GPIO 15) and SD CS (GPIO 5) are never LOW at the same time — Arduino handles this automatically.'
))

# 15.2 SPI Bus sharing diagram
new_elems.append(make_sub_heading_para('15.2  SPI Bus Sharing Diagram'))
for p in code_paras(
"""ESP32 VSPI Bus
      │
      ├── GPIO 15 CS ──► ILI9486 Display
      ├── GPIO 5  CS ──► MicroSD Card
      │
      └── GPIO 18 CLK  ─┐
          GPIO 19 MISO ─┤── Shared by both devices
          GPIO 23 MOSI ─┘

Only one CS pin is LOW at any time → no bus conflict"""
):
    new_elems.append(p)

# 15.3 SD File Structure
new_elems.append(make_sub_heading_para('15.3  SD Card File Structure'))
for p in code_paras(
"""/voicedesk/
  ├── config.json            ← WiFi credentials, Pi5 server IP
  ├── queue/
  │     ├── task_1717401600001.json     ← pending items (not yet synced)
  │     ├── meeting_1717401600002.json
  │     └── ...
  └── synced/
        └── (files moved here after successful Pi5 sync)"""
):
    new_elems.append(p)

new_elems.append(new_para(
    'Each queue file is a self-contained JSON record with all fields needed to reconstruct '
    'the item on the Pi5 (type, title, sheet, date, remarks, timestamp). Files are named '
    'with a millisecond timestamp to avoid collisions.',
    size=10, color=C_DARK
))

# 15.4 Sync State Machine
new_elems.append(make_sub_heading_para('15.4  Sync State Machine'))
for p in code_paras(
"""On every write attempt:
  ├── WiFi connected + Pi5 reachable?
  │     YES ──► Send directly to Pi5 via WebSocket
  │              SD not used. Sheets sync runs on Pi5.
  │
  │     NO  ──► Write JSON file to /voicedesk/queue/
  │              Display: 🔴 [Offline — 3 queued]

On WiFi reconnect event (auto):
  └── Call syncQueueToPi5()
       ├── Read each file in /voicedesk/queue/
       ├── POST to Pi5 /api/offline_batch
       ├── Pi5 writes to SQLite + Google Sheets
       └── On success: move file to /voicedesk/synced/
            Display: ✅ [Synced 3 items]

Manual sync (two ways):
  ├── Hold OK button 2 seconds  →  triggers syncQueueToPi5()
  └── Voice: "sync my data"     →  Pi5 receives SYNC_NOW intent"""
):
    new_elems.append(p)

# 15.5 Arduino Code
new_elems.append(make_sub_heading_para('15.5  Arduino Code Reference'))

new_elems.append(new_para('Initialisation (both SPI devices):', size=10, color=C_DARK))
for p in code_paras(
"""#include <SPI.h>
#include <SD.h>
#include <TFT_eSPI.h>

#define SD_CS   5      // SD chip select
#define TFT_CS  15     // Display chip select (in TFT_eSPI User_Setup.h)

TFT_eSPI tft;

void setup() {
    tft.begin();           // TFT_eSPI takes TFT_CS
    if (!SD.begin(SD_CS)) {
        Serial.println("SD init failed");
    }
    SD.mkdir("/voicedesk/queue");
    SD.mkdir("/voicedesk/synced");
}"""
):
    new_elems.append(p)

new_elems.append(new_para('Write to offline queue:', size=10, color=C_DARK))
for p in code_paras(
"""void queueOffline(String type, String jsonPayload) {
    String path = "/voicedesk/queue/" + type + "_"
                + String(millis()) + ".json";
    File f = SD.open(path, FILE_WRITE);
    if (f) { f.println(jsonPayload); f.close(); }
    updateDisplay(OFFLINE_QUEUED);  // show red indicator
}"""
):
    new_elems.append(p)

new_elems.append(new_para('Auto-sync on WiFi reconnect:', size=10, color=C_DARK))
for p in code_paras(
"""void syncQueueToPi5() {
    File dir = SD.open("/voicedesk/queue/");
    int count = 0, ok = 0;
    while (true) {
        File entry = dir.openNextFile();
        if (!entry) break;
        count++;
        String payload = "";
        while (entry.available()) payload += (char)entry.read();
        String name = entry.name();
        entry.close();

        if (postToPi5("/api/offline_batch", payload)) {
            SD.rename(("/voicedesk/queue/" + name).c_str(),
                      ("/voicedesk/synced/" + name).c_str());
            ok++;
        }
    }
    dir.close();
    showSyncResult(ok, count);  // "Synced 3/3 items" on display
}

// Register auto-sync on WiFi connect:
WiFi.onEvent([](WiFiEvent_t e, WiFiEventInfo_t info) {
    if (e == ARDUINO_EVENT_WIFI_STA_GOT_IP)
        syncQueueToPi5();
});"""
):
    new_elems.append(p)

new_elems.append(new_para('Manual sync (button hold):', size=10, color=C_DARK))
for p in code_paras(
"""// In button handler loop:
if (btn_ok_held_ms >= 2000) {
    showScreen(SCREEN_SYNCING);
    syncQueueToPi5();
}"""
):
    new_elems.append(p)

# 15.6 Display Indicators
new_elems.append(make_sub_heading_para('15.6  Display Status Indicators'))

tmp2 = Document()
ind_tbl = make_table_xml(tmp2,
    ['State', 'Home Screen Indicator', 'Colour'],
    [
        ('Online, synced',         '[WiFi ✓]  [Sync ✓]',            'Green'),
        ('Online, queue pending',  '[WiFi ✓]  [Syncing... 3]',       'Amber (auto-sync running)'),
        ('Offline, queuing',       '[WiFi ✗]  [SD 3 pending]',       'Red'),
        ('Manual sync done',       '[Sync ✓  3 items sent]',         'Green flash then normal'),
        ('SD card missing/error',  '[SD ✗  — local only]',           'Orange (degraded mode)'),
    ],
    col_widths=[1.8, 2.4, 1.6]
)
new_elems.append(ind_tbl)
new_elems.append(new_para('', size=10, space_before=0, space_after=4))

new_elems.append(make_note_para(
    'If the SD card is missing or corrupt, the device still works — voice commands are sent '
    'directly to Pi5 when online. Offline mode is simply unavailable. The display shows '
    '[SD ✗] to indicate degraded mode.'
))

# ── Insert all before "— End of Document —"
insert_before(body, end_para._p, *new_elems)
print('✓ Section 15 inserted')

# ══════════════════════════════════════════════════════════════════════════════
# 10. Update meta table status
# ══════════════════════════════════════════════════════════════════════════════
for tbl in doc.tables[:1]:  # first table = meta table
    for row in tbl.rows:
        if 'Status' in row.cells[0].text:
            for run in row.cells[1].paragraphs[0].runs:
                run.text = 'Architecture v1.1 — SD Card Offline Storage Added'
    break
print('✓ Meta status updated')

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/sessions/sharp-affectionate-dijkstra/mnt/outputs/VoiceDesk_Architecture.docx'
doc.save(out)
print(f'\n✅ Saved: {out}')
